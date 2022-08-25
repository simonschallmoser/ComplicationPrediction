import numpy as np
import pandas as pd
import sys
sys.path.insert(0, '../prediction')

from docx import Document
from docx.shared import Cm

import prediction
import data_reader
import helper
import preprocess
import data_imputation
from categories_dicts import ordered_tests, rename_tests, drug_class_to_name
from icd9 import icd9_to_description

def demo_characteristics(data, outcomes, ids_population_all, ids_outcome_all, baseline=2008):
    demo_dict = {}
    # Characteristics of all patients at baseline
    data_baseline = data.loc[pd.IndexSlice[:, baseline], :]
    demo_dict['female_base'] = str(int(data_baseline.sex.sum()))
    demo_dict['female_base_pct'] = str(np.round(data_baseline.sex.sum() / len(data_baseline) * 100, 1))
    demo_dict['male_base'] = int(len(data_baseline) - data_baseline.sex.sum())
    demo_dict['male_base_pct'] = str(np.round(demo_dict['male_base'] / len(data_baseline) * 100, 1))
    cols = ['age', 'bmi', 'systolic_bp', 'diastolic_bp']
    for col in cols:
        demo_dict[col+'_base_mean'] = str(np.round(data_baseline[col].mean(), 1))
        demo_dict[col+'_base_std'] = str(np.round(data_baseline[col].std(), 1))
    for outcome in outcomes:
        ids_population = ids_population_all[outcome]
        ids_outcome = ids_outcome_all[outcome]
        data_outcome_baseline = data_baseline.loc[ids_outcome]
        data_outcome_diagnosis = data.loc[[(k, ids_outcome[k]) for k in ids_outcome.keys()]]
        # Sex
        demo_dict['female_'+outcome] = str(int(data_outcome_baseline.sex.sum()))
        demo_dict['female_'+outcome+'_pct'] = str(np.round(data_outcome_baseline.sex.sum() 
                                                       / len(data_outcome_baseline) * 100, 1))
        demo_dict['male_'+outcome] = int(len(data_outcome_baseline) - data_outcome_baseline.sex.sum())
        demo_dict['male_'+outcome+'_pct'] = str(np.round(demo_dict['male_'+outcome] 
                                                       / len(data_outcome_baseline) * 100, 1))
        # Other predictors
        for col in cols:
            demo_dict[col+'_base_'+outcome+'_mean'] = str(np.round(data_outcome_baseline[col].mean(), 1))
            demo_dict[col+'_base_'+outcome+'_std'] = str(np.round(data_outcome_baseline[col].std(), 1))
            demo_dict[col+'_diag_'+outcome+'_mean'] = str(np.round(data_outcome_diagnosis[col].mean(), 1))
            demo_dict[col+'_diag_'+outcome+'_std'] = str(np.round(data_outcome_diagnosis[col].std(), 1))
            
    return demo_dict

def add_hba1c_mmol_mol(data):
    data_new = data.copy()
    data_hba1c = data_new['test=hba1c(%)']
    data_hba1c = data_hba1c.round(1)
    for i in np.arange(3.0, 18.0, 0.1):
        data_hba1c.replace(i, np.round(10.93 * i - 23.5, 0), inplace=True)
    data_new['test=hba1c(mmol/mol)_1'] = data_hba1c
    return data_new

def lab_characteristics(data, lab_cols, outcomes, ids_outcome_all, baseline=2008):
    lab_dict = {}
    data_baseline = data.loc[pd.IndexSlice[:, baseline], :]
    for col in lab_cols:
        lab_dict[col+'_base_mean'] = str(np.round(data_baseline[col].mean(), 2))
        lab_dict[col+'_base_std'] = str(np.round(data_baseline[col].std(), 2))
    for outcome in outcomes:
        data_baseline_outcome = data_baseline.loc[ids_outcome_all[outcome]]
        data_diagnosis = data.loc[[(k, ids_outcome_all[outcome][k]) for k in ids_outcome_all[outcome].keys()]]
        for col in lab_cols:
            lab_dict[col+'_base_'+outcome+'_mean'] = str(np.round(data_baseline_outcome[col].mean(), 2))
            lab_dict[col+'_base_'+outcome+'_std'] = str(np.round(data_baseline_outcome[col].std(), 2))
            lab_dict[col+'_diag_'+outcome+'_mean'] = str(np.round(data_diagnosis[col].mean(), 2))
            lab_dict[col+'_diag_'+outcome+'_std'] = str(np.round(data_diagnosis[col].std(), 2))
    
    return lab_dict

def icd_characteristics(data, icd_codes, outcomes, ids_outcome_all, baseline=2008):
    icd_dict = {}
    data_baseline = data.loc[pd.IndexSlice[:, baseline], :]
    for code in icd_codes:
        icd_dict[code+'_base'] = str(int(data_baseline[code].sum()))
        icd_dict[code+'_base_pct'] = str(np.round(data_baseline[code].sum() / len(data_baseline) * 100, 1))
    for outcome in outcomes:
        data_baseline_outcome = data_baseline.loc[ids_outcome_all[outcome]]
        data_diagnosis = data.loc[[(k, ids_outcome_all[outcome][k]) for k in ids_outcome_all[outcome].keys()]]
        for code in icd_codes:
            icd_dict[code+'_base_'+outcome] = str(int(data_baseline_outcome[code].sum()))
            icd_dict[code+'_base_'+outcome+'_pct'] = str(np.round(data_baseline_outcome[code].sum() 
                                                                 / len(data_baseline_outcome) * 100, 1))
            icd_dict[code+'_diag_'+outcome] = str(int(data_diagnosis[code].sum()))
            icd_dict[code+'_diag_'+outcome+'_pct'] = str(np.round(data_diagnosis[code].sum() 
                                                                 / len(data_diagnosis) * 100, 1))
    return icd_dict

def med_characteristics(data, outcomes, ids_outcome_all, baseline=2008):
    med_dict = {}
    data_baseline = data.loc[pd.IndexSlice[:, baseline], :]
    med_cols = [i for i in data_baseline.columns if 'med' in i]
    for col in med_cols:
        med_dict[col+'_base'] = str(int(data_baseline[col].sum()))
        med_dict[col+'_base_pct'] = str(np.round(data_baseline[col].sum() / len(data_baseline) * 100, 1))
    for outcome in outcomes:
        data_baseline_outcome = data_baseline.loc[ids_outcome_all[outcome]]
        data_diagnosis = data.loc[[(k, ids_outcome_all[outcome][k]) for k in ids_outcome_all[outcome].keys()]]
        for col in med_cols:
            med_dict[col+'_base_'+outcome] = str(int(data_baseline_outcome[col].sum()))
            med_dict[col+'_base_'+outcome+'_pct'] = str(np.round(data_baseline_outcome[col].sum() 
                                                     / len(data_baseline_outcome) * 100, 1))
            med_dict[col+'_diag_'+outcome] = str(int(data_diagnosis[col].sum()))
            med_dict[col+'_diag_'+outcome+'_pct'] = str(np.round(data_diagnosis[col].sum() 
                                                     / len(data_diagnosis) * 100, 1))
    return med_dict

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def table_patient_characteristics(outcomes, outcomes_long, population, evaluation,
                                  data, ids_population_all, ids_outcome_all,
                                  baseline, n_tests, n_icd9_codes, format_,
                                  add_diagnosis_criteria=False,
                                  style='TableGrid', col_width=4, 
                                  table_name='patient_characteristics', both=False, p_values=False):
    
    # Adjust table name if p-values are calculated
    if p_values:
        table_name = 'patient_characteristics_p_values'

    # Initizalize table
    
    doc = Document()
    table = doc.add_table(0, 0)
    table.style = style
    
    if both:
        for i in range(14):
            table.add_column(Cm(col_width))
        if p_values:
            table.add_column(Cm(col_width))
    else:
        for i in range(8):
            table.add_column(Cm(col_width))
        if p_values:
            table.add_column(Cm(col_width))
            
    table.add_row()
    row = table.rows[0]
    row.cells[1].text = 'Characteristics of patients at baseline'
    if both:
        for idx, outcome in enumerate(outcomes_long):
            row.cells[2*idx+2].text = outcome
            if p_values:
                row.cells[-1].text = 'p-value'
    else:
        for idx, outcome in enumerate(outcomes_long):
            row.cells[idx+2].text = outcome
            if p_values:
                row.cells[-1].text = 'p-value'
    if add_diagnosis_criteria:        
        # Add diagnosis criteria
        row_length = len(table.rows)
        table.add_row()
        row = table.rows[1]
        row.cells[0].text = 'Diagnosis criteria'
        make_rows_bold

        row_length = len(table.rows)
        table.add_row()
        row = table.rows[row_length]
        row.cells[0].text = '    ICD-9 codes'
        for idx, outcome in enumerate(outcomes):
            row.cells[idx+2].text = str(prediction.outcome_code(outcome))

        row_length = len(table.rows)
        table.add_row()
        row = table.rows[row_length]
        row.cells[0].text = '    Biomarkers'
        for idx, outcome in enumerate(outcomes):
            if outcome == 'renal':
                row.cells[idx+2].text = 'eGFR<60 [ml/min], one measurement; \
                                         albumin to creatinine in the urine ≥30, two measurements'
            else:
                row.cells[idx+2].text = '-'
     
    # Sample size characteristics
    row_length = len(table.rows)
    table.add_row()
    row_n = table.rows[row_length]
    row_n.cells[0].text = 'Number of patients'
    row_n.cells[1].text = str(len(data.index.get_level_values('id').unique()))
    for idx, outcome in enumerate(outcomes):
        row_n.cells[idx+2].text = str(len(ids_outcome_all[outcome]))
    row_length = len(table.rows)
    table.add_row()
    row = table.rows[row_length]
    row.cells[0].text = 'Number of patients without complication at baseline'
    row.cells[1].text = '-'
    for idx, outcome in enumerate(outcomes):
        row.cells[idx+2].text = str(len(ids_population_all[outcome]))
    row_length = len(table.rows)
    table.add_row()
    row_incidence = table.rows[row_length]
    row_incidence.cells[0].text = '5-year incidence [%]'
    row_incidence.cells[1].text = '-'
    for idx, outcome in enumerate(outcomes):
        row_incidence.cells[idx+2].text = str(np.round(len(ids_outcome_all[outcome]) 
                                                       / len(ids_population_all[outcome]) * 100, 1))
    
    # Demo characteristics
    demo_dict = demo_characteristics(data, outcomes, ids_population_all, ids_outcome_all)
    row_length = len(table.rows)
    table.add_row()
    row = table.rows[row_length]
    row.cells[0].text = 'Demographic data'
    make_rows_bold(row)
    row_length = len(table.rows)
    table.add_row()
    row = table.rows[row_length]
    row.cells[0].text = 'Sex'
    row_length = len(table.rows)
    table.add_row()
    row = table.rows[row_length]
    row.cells[0].text = '    Female'
    row.cells[1].text = demo_dict['female_base'] + ' (' + demo_dict['female_base_pct'] + '%)'
    for idx, outcome in enumerate(outcomes):
        row.cells[idx+2].text = demo_dict['female_'+outcome] + ' (' + demo_dict['female_'+outcome+'_pct'] + '%)'
    row_length = len(table.rows)
    table.add_row()
    row = table.rows[row_length]
    row.cells[0].text = '    Male'
    row.cells[1].text = str(demo_dict['male_base']) + ' (' + demo_dict['male_base_pct'] + '%)'
    for idx, outcome in enumerate(outcomes):
        row.cells[idx+2].text = str(demo_dict['male_'+outcome]) + ' (' + demo_dict['male_'+outcome+'_pct'] + '%)'
    demo_cols = ['age', 'bmi', 'systolic_bp', 'diastolic_bp']
    demo_cols_table = ['Age [years]', 'BMI [kg/m^2]', 
                       'Systolic blood pressure [mmHg]', 'Diastolic blood pressure [mmHg]']
    for col, col_table in zip(demo_cols, demo_cols_table):
        row_length = len(table.rows)
        table.add_row()
        row = table.rows[row_length]
        row.cells[0].text = col_table
        row.cells[1].text = demo_dict[col+'_base_mean'] + ' (' + demo_dict[col+'_base_std'] + ')'
        for idx, outcome in enumerate(outcomes):
            if evaluation == 'baseline':
                row.cells[idx+2].text = demo_dict[col+'_base_'+outcome+'_mean'] + ' (' \
                                        + demo_dict[col+'_base_'+outcome+'_std'] + ')'
            elif evaluation == 'diagnosis':
                row.cells[idx+2].text = demo_dict[col+'_diag_'+outcome+'_mean'] + ' (' \
                                        + demo_dict[col+'_diag_'+outcome+'_std'] + ')'
            else:
                raise ValueError('Unknown evaluation parameter.')
    
    # Lab characteristics
    data_baseline = data.loc[pd.IndexSlice[:, baseline], :]
    test_columns = [i for i in data_baseline.columns if 'test' in i]
    lab_cols = list(data_baseline[test_columns].isna().sum().sort_values(ascending=True)[:n_tests].index)
    lab_cols.append('test=hba1c(mmol/mol)_1')
    lab_dict = lab_characteristics(add_hba1c_mmol_mol(data), lab_cols, outcomes, ids_outcome_all)
    
    row_length = len(table.rows)
    table.add_row()
    row = table.rows[row_length]
    row.cells[0].text = 'Biomarkers'
    make_rows_bold(row)
    row_length = len(table.rows)
    ordered_tests_1 = [i for i in ordered_tests if i in lab_cols]
    if format_ == 'short':
        ordered_tests_1 = ['test=glucose', 'test=hba1c(%)', 'test=hba1c(mmol/mol)_1']
    for row_idx, test in enumerate(ordered_tests_1):
        table.add_row()
        row = table.rows[row_length+row_idx]
        row.cells[0].text = rename_tests[test]
        row.cells[1].text = lab_dict[test+'_base_mean'] + ' (' + lab_dict[test+'_base_std'] + ')'
        for idx, outcome in enumerate(outcomes):
            if evaluation == 'baseline':
                row.cells[idx+2].text = lab_dict[test+'_base_'+outcome+'_mean'] + ' (' + \
                                        lab_dict[test+'_base_'+outcome+'_std'] + ')'
            elif evaluation == 'diagnosis':
                row.cells[idx+2].text = lab_dict[test+'_diag_'+outcome+'_mean'] + ' (' + \
                                        lab_dict[test+'_diag_'+outcome+'_std'] + ')'
            else:
                raise ValueError('Unknown evaluation parameter.')
                
    # ICD-9 characteristics
    if format_ != 'short':
        icd9_columns = [i for i in data_baseline.columns if 'icd9_prefix' in i]
        icd9_columns = data_baseline[icd9_columns].sum().sort_values(ascending=True)[-n_icd9_codes:].index
        icd_dict = icd_characteristics(data, icd9_columns, outcomes, ids_outcome_all)

        row_length = len(table.rows)
        table.add_row()
        row = table.rows[row_length]
        row.cells[0].text = 'Disease codes'
        make_rows_bold(row)
        row_length = len(table.rows)
        if format_ == 'short':
            icd9_columns = ['icd9_prefix=401', 'icd9_prefix=272']
        for row_idx, icd9_column in enumerate(icd9_columns):
            table.add_row()
            row = table.rows[row_length+row_idx]
            code = int(icd9_column.split('=')[-1])
            row.cells[0].text = icd9_to_description[code]
            if code == 401:
                row.cells[0].text = 'Hypertension'
            if code == 272:
                row.cells[0].text = 'Dyslipidemia'
            row.cells[1].text = icd_dict[icd9_column+'_base'] + ' (' + icd_dict[icd9_column+'_base_pct'] + '%)'
            for idx, outcome in enumerate(outcomes):
                if evaluation == 'baseline':
                    row.cells[idx+2].text = icd_dict[icd9_column+'_base_'+outcome+''] + ' (' + \
                                            icd_dict[icd9_column+'_base_'+outcome+'_pct'] + '%)'
                elif evaluation == 'diagnosis':
                    row.cells[idx+2].text = icd_dict[icd9_column+'_diag_'+outcome+''] + ' (' + \
                                            icd_dict[icd9_column+'_diag_'+outcome+'_pct'] + '%)'
                else:
                    raise ValueError('Unknown evaluation parameter.')

    # Med characteristics
    if format_ != 'short':
        med_dict = med_characteristics(data, outcomes, ids_outcome_all)
        row_length = len(table.rows)
        table.add_row()
        row = table.rows[row_length]
        row.cells[0].text = 'Medications'
        make_rows_bold(row)
        row_length = len(table.rows)    
        med_cols = [i for i in data.columns if 'med' in i]
        for row_idx, med_col in enumerate(med_cols):
            table.add_row()
            row = table.rows[row_length+row_idx]
            med_class = int(med_col.split('= ')[-1].split(')')[0])
            row.cells[0].text = drug_class_to_name[med_class]
            if med_class == 3:
                row.cells[0].text = 'Other antihypertensive drug'
            row.cells[1].text = med_dict[med_col+'_base'] + ' (' + med_dict[med_col+'_base_pct'] + '%)'
            for idx, outcome in enumerate(outcomes):
                if evaluation == 'baseline':
                    row.cells[idx+2].text = med_dict[med_col+'_base_'+outcome+''] + ' (' + \
                                            med_dict[med_col+'_base_'+outcome+'_pct'] + '%)'
                elif evaluation == 'diagnosis':
                    row.cells[idx+2].text = med_dict[med_col+'_diag_'+outcome+''] + ' (' + \
                                            med_dict[med_col+'_diag_'+outcome+'_pct'] + '%)'
                else:
                    raise ValueError('Unknown evaluation parameter.')
    
    doc.save(f'{table_name}_{population}_{format_}_{evaluation}.docx')